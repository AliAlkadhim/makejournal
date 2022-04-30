import datetime as dt
import sys
import pip
def install(package):
    pip.main(['install', package])
try:
    import docx
except ImportError:
    print('docx is not installed, installing it now!')
    install('docx')

from docx import Document
import argparse

today = dt.date.today()

parser = argparse.ArgumentParser(description='Generate journal')

parser.add_argument('--year', type=int,default=today.year,  required=False, help = '''year, default is the current year, 2022,
''')
parser.add_argument('--start_month', type=int, required=True, help = '''start date month
''')
parser.add_argument('--start_day', type=int, required=True, help = '''start date day of the month
''')
parser.add_argument('--end_month', type=int, required=True, help = '''end date month
''')
parser.add_argument('--end_day', type=int, required=True, help = '''end date day of the month
''')



args = parser.parse_args()

# The date is format YYYY-MM-DD
start_date=dt.date(args.year, args.start_month, args.start_day) 

# duration = dt.timedelta(days=66)
end_date = dt.date(args.year,args.end_month, args.end_day) 
duration = end_date - start_date

name = 'JOURNAL_STARTING_FROM_%s_TO_%s'  % (start_date, end_date)
def make_journal():
    """Usage: python makejournal.py --start_month 5 --start_day 12"""
    document = Document()
    document.add_heading('YOU HAVE %d DAYS IN THIS JOURNAL, MAKE EVERY SECOND COUNT!' %duration.days, 0)
    # document.add_heading('My Stoic Journal', 0)

    for i in range(duration.days+1):
        day = start_date+dt.timedelta(days=i)
        document.add_heading(str(day), level=0)
        a1=document.add_paragraph('What did you do well today?', 
                    style='List Bullet')
        a1.add_run('\n\n\n')
        a2=document.add_paragraph('What did you do not so well today?', 
                    style='List Bullet')
        a2.add_run('\n\n\n')			  
        a3=document.add_paragraph('What will you do tomorrow to improve?', 
                    style='List Bullet')
        a3.add_run('\n\n\n')
        a4 = document.add_paragraph('RANDOM NOTES',)
        document.add_page_break()
    document.save(name+'.docx')

if __name__ == '__main__':
    # print(name)
    make_journal()